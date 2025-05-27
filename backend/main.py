import os
import re
import hashlib
import tempfile
import json
import concurrent.futures
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import pickle
import matplotlib.pyplot as plt
import networkx as nx
import logging
import traceback
from datetime import datetime
from logging.handlers import RotatingFileHandler
from functools import lru_cache
import signal
import sys

# FastAPI imports
from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import uvicorn

# Import routers
from routers.documents import router as documents_router
from routers.chat import router as chat_router
from routers.notes import router as notes_router
from routers.flashcards import router as flashcards_router
from routers.mindmaps import router as mindmaps_router
from routers.tests import router as tests_router
from routers.roadmaps import router as roadmaps_router
from routers.dsa import router as dsa_router
from routers.progress import router as progress_router
from routers.forum import router as forum_router  # Import forum router
from routers.exam_practice import router as exam_practice_router  # Import exam practice router

# Document processing
import fitz  # PyMuPDF
import PyPDF2
from docx import Document

# Language model and embeddings
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

# Import models and utilities
from models.schemas import (
    DocumentMetadata, DocumentResponse, 
    ChatMessage, ChatRequest, ChatResponse,
    MindMapResponse, FlashcardResponse, TestResponse, RoadmapResponse,
    ErrorResponse
)
from agents import (
    create_study_tutor_agent,
    create_explanation_task,
    create_note_taker_agent,
    create_assessment_expert_agent,
    create_flashcard_specialist_agent,
    create_visual_learning_expert_agent,
    create_learning_coach_agent,
    create_roadmap_planner_agent,
    create_question_fetching_agent,
    create_filtering_agent,
    create_progress_tracking_agent,
    create_personalization_agent,
    create_debugging_agent,
    create_dsa_recommendation_agent,
    create_coding_pattern_agent,
    create_interview_strategy_agent,
    create_company_specific_agent,    create_notes_generation_task,
    create_test_generation_task,
    create_enhanced_test_generation_task,
    create_flashcard_generation_task,
    create_mind_map_task,
    create_progress_analysis_task,
    create_roadmap_generation_task,
    create_quick_roadmap_generation_task,
    create_question_fetching_task,
    create_filtering_task,
    create_progress_tracking_task,
    create_personalization_task,
    create_debugging_task,
    create_dsa_recommendation_task,
    create_pattern_identification_task,
    create_company_preparation_task,
    create_mock_interview_task,
    run_agent_task
)

# Configuration constants
MAX_CHUNK_SIZE = 1000
OVERLAP_SIZE = 200
BATCH_SIZE = 10  # Number of pages to process at once for memory management
LOG_FORMAT = "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
LOG_FILE = "backend/logs/app.log"
MAX_LOG_SIZE = 10 * 1024 * 1024  # 10MB
BACKUP_COUNT = 5

# Custom exception classes
class DocumentProcessingError(Exception):
    """Base exception for document processing errors"""
    pass

class FileFormatError(DocumentProcessingError):
    """Exception for unsupported file formats"""
    pass

class TextExtractionError(DocumentProcessingError):
    """Exception for text extraction failures"""
    pass

class VectorStoreError(Exception):
    """Exception for vector store operations"""
    pass

class AgentError(Exception):
    """Exception for agent-related operations"""
    pass

# Setup logging
def setup_logging():
    """Configure logging with file rotation and console output"""
    # Create logs directory if it doesn't exist
    log_dir = Path("backend/logs")
    log_dir.mkdir(parents=True, exist_ok=True)
    
    # Create formatter
    formatter = logging.Formatter(LOG_FORMAT)
    
    # Setup file handler with rotation
    file_handler = RotatingFileHandler(
        LOG_FILE,
        maxBytes=MAX_LOG_SIZE,
        backupCount=BACKUP_COUNT
    )
    file_handler.setFormatter(formatter)
    
    # Setup console handler
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)
    
    # Setup root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)
    
    # Create logger for this module
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.DEBUG)
    
    # Reduce noise from watchfiles module
    logging.getLogger('watchfiles').setLevel(logging.ERROR)
    logging.getLogger('watchgod').setLevel(logging.ERROR)
    
    return logger

logger = setup_logging()

# Ensure storage directories exist
storage_dirs = [
    "./storage",
    "./storage/documents",
    "./storage/vectorstores",
    "./storage/cache",
    "./storage/notes",
    "./storage/flashcards",
    "./storage/tests",
    "./storage/mindmaps",
    "./storage/roadmaps"
]

for dir_path in storage_dirs:
    Path(dir_path).mkdir(parents=True, exist_ok=True)
    logger.info(f"Ensured directory exists: {dir_path}")

# Initialize FastAPI app
app = FastAPI(
    title="InterviewBuddy API",
    description="API for Interview Preparation and Study Assistant",
    version="1.0.0"
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:8000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Register routers with prefix
app.include_router(documents_router, prefix="/api/documents", tags=["documents"])
app.include_router(chat_router, prefix="/api/chat", tags=["chat"])
app.include_router(notes_router, prefix="/api/notes", tags=["notes"])
app.include_router(flashcards_router, prefix="/api/flashcards", tags=["flashcards"])
app.include_router(mindmaps_router, prefix="/api/mindmaps", tags=["mindmaps"])
app.include_router(tests_router, prefix="/api/tests", tags=["tests"])
app.include_router(roadmaps_router, prefix="/api/roadmaps", tags=["roadmaps"])
app.include_router(dsa_router, prefix="/api/dsa", tags=["dsa"])
app.include_router(progress_router, prefix="/api/progress", tags=["progress"])
app.include_router(forum_router, prefix="/api/forum", tags=["forum"])  # Register forum router
app.include_router(exam_practice_router, prefix="/api/exam-practice", tags=["exam-practice"])  # Register exam practice router

# Error handling middleware
@app.middleware("http")
async def error_handling_middleware(request: Request, call_next):
    """Global error handling middleware"""
    try:
        return await call_next(request)
    except Exception as e:
        # Log the full error with traceback
        logger.error(
            f"Unhandled error in {request.method} {request.url.path}",
            exc_info=True,
            extra={
                "error_type": type(e).__name__,
                "error_message": str(e),
                "path": request.url.path,
                "method": request.method
            }
        )
        
        # Return appropriate error response
        if isinstance(e, HTTPException):
            return JSONResponse(
                status_code=e.status_code,
                content={"detail": e.detail}
            )
        elif isinstance(e, DocumentProcessingError):
            return JSONResponse(
                status_code=400,
                content={"detail": str(e)}
            )
        elif isinstance(e, VectorStoreError):
            return JSONResponse(
                status_code=500,
                content={"detail": f"Vector store error: {str(e)}"}
            )
        elif isinstance(e, AgentError):
            return JSONResponse(
                status_code=500,
                content={"detail": f"Agent error: {str(e)}"}
            )
        else:
            return JSONResponse(
                status_code=500,
                content={"detail": "Internal server error"}
            )

# Utility Functions
def get_embeddings():
    """Get cached instance of HuggingFace embeddings"""
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2"
    )

def compute_file_hash(file_bytes):
    """Compute MD5 hash for a file to use as cache key"""
    try:
        return hashlib.md5(file_bytes).hexdigest()
    except Exception as e:
        logger.error("Failed to compute file hash", exc_info=True)
        raise DocumentProcessingError("Failed to process file") from e

def extract_text_from_pdf_pymupdf(file_bytes, start_page=0, end_page=None):
    """Extract text from PDF using PyMuPDF (primary method)"""
    try:
        text_by_page = []
        
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
            total_pages = pdf_document.page_count
            end_page = end_page if end_page is not None else total_pages
            
            for page_num in range(start_page, min(end_page, total_pages)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                text_by_page.append(text)
                
        return text_by_page
    except Exception as e:
        logger.error("PyMuPDF text extraction failed", exc_info=True)
        raise TextExtractionError("Failed to extract text from PDF using PyMuPDF") from e

def extract_text_from_pdf_pypdf2(file_bytes, start_page=0, end_page=None):
    """Extract text from PDF using PyPDF2 (fallback method)"""
    try:
        text_by_page = []
        
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_file_path = temp_file.name
        
        try:
            with open(temp_file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                end_page = end_page if end_page is not None else total_pages
                
                for page_num in range(start_page, min(end_page, total_pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_by_page.append(text)
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
        return text_by_page
    except Exception as e:
        logger.error("PyPDF2 text extraction failed", exc_info=True)
        raise TextExtractionError("Failed to extract text from PDF using PyPDF2") from e

def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file"""
    try:
        text_by_page = []
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(file_bytes)
            temp_file_path = temp_file.name
        
        try:
            doc = Document(temp_file_path)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Split into pseudo-pages (approx. 3000 chars per page)
            chars_per_page = 3000
            text_by_page = [full_text[i:i+chars_per_page] 
                           for i in range(0, len(full_text), chars_per_page)]
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
        return text_by_page
    except Exception as e:
        logger.error("DOCX text extraction failed", exc_info=True)
        raise TextExtractionError("Failed to extract text from DOCX file") from e

def extract_text_from_txt(file_bytes):
    """Extract text from TXT file"""
    try:
        text = file_bytes.decode('utf-8')
        
        # Split into pseudo-pages (approx. 3000 chars per page)
        chars_per_page = 3000
        text_by_page = [text[i:i+chars_per_page] 
                       for i in range(0, len(text), chars_per_page)]
        
        return text_by_page
    except Exception as e:
        logger.error("TXT text extraction failed", exc_info=True)
        raise TextExtractionError("Failed to extract text from TXT file") from e

async def process_document_async(file: UploadFile):
    """Process document and create vector store"""
    try:
        # Read file bytes
        file_bytes = await file.read()
        file_hash = compute_file_hash(file_bytes)
        file_extension = file.filename.split('.')[-1].lower()
        
        logger.info(f"Processing document: {file.filename} ({file_extension})")
        
        # Check if document is already cached
        cache_path = Path("./storage/cache") / f"{file_hash}.pkl"
        if cache_path.exists():
            logger.info(f"Found cached document: {file_hash}")
            with open(cache_path, 'rb') as f:
                return pickle.load(f)
        
        # Extract text based on file type
        if file_extension == 'pdf':
            try:
                # Use PyMuPDF first
                text_by_page = extract_text_from_pdf_pymupdf(file_bytes)
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed, falling back to PyPDF2: {e}")
                text_by_page = extract_text_from_pdf_pypdf2(file_bytes)
        elif file_extension == 'docx':
            text_by_page = extract_text_from_docx(file_bytes)
        elif file_extension == 'txt':
            text_by_page = extract_text_from_txt(file_bytes)
        else:
            raise FileFormatError(f"Unsupported file format: {file_extension}")
        
        # Create document object
        doc_info = {
            'id': file_hash,
            'filename': file.filename,
            'pages': len(text_by_page),
            'text_by_page': text_by_page,
            'file_hash': file_hash,
            'created_at': datetime.now().isoformat(),
            'file_size': len(file_bytes),
            'mime_type': file.content_type
        }
        
        # Cache the processed document
        with open(cache_path, 'wb') as f:
            pickle.dump(doc_info, f)
        
        logger.info(f"Successfully processed document: {file.filename}")
        return doc_info
    except Exception as e:
        logger.error(f"Error processing document: {file.filename}", exc_info=True)
        if isinstance(e, (FileFormatError, TextExtractionError)):
            raise
        raise DocumentProcessingError(f"Failed to process document: {str(e)}") from e

def create_vector_store(document_info: dict) -> FAISS:
    """Create a vector store from document text"""
    try:
        # Combine all pages into a single text
        full_text = "\n\n".join(document_info['text_by_page'])
        
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=MAX_CHUNK_SIZE,
            chunk_overlap=OVERLAP_SIZE,
            length_function=len,
        )
        chunks = text_splitter.split_text(full_text)
        
        # Create documents for vector store
        documents = [
            LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": document_info['filename'],
                    "page": i // (MAX_CHUNK_SIZE // OVERLAP_SIZE) + 1
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # Create and return vector store
        embeddings = get_embeddings()
        vector_store = FAISS.from_documents(documents, embeddings)
        return vector_store
    except Exception as e:
        logger.error(f"Error creating vector store: {e}")
        raise VectorStoreError(f"Error creating vector store: {str(e)}")

def get_relevant_context(vector_store, query: str, k: int = 3) -> str:
    """Get relevant context from vector store for a query"""
    try:
        docs = vector_store.similarity_search(query, k=k)
        return "\n\n".join([doc.page_content for doc in docs])
    except Exception as e:
        logger.error(f"Error getting relevant context: {e}")
        raise VectorStoreError(f"Error getting relevant context: {str(e)}")

# Health Check Endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        # Check storage directories
        for dir_path in storage_dirs:
            if not Path(dir_path).exists():
                raise Exception(f"Storage directory not found: {dir_path}")
        
        # Check if we can create embeddings
        embeddings = get_embeddings()
        
        logger.info("Health check passed")
        return {"status": "healthy", "message": "API is running"}
    except Exception as e:
        logger.error("Health check failed", exc_info=True)
        return {"status": "unhealthy", "message": str(e)}

# Also add the /api/health endpoint for consistency
@app.get("/api/health")
async def api_health_check():
    """API health check endpoint"""
    return await health_check()

def handle_shutdown(signum, frame):
    """Handle shutdown signals gracefully"""
    logger.info(f"Received shutdown signal {signum}")
    sys.exit(0)

# Register signal handlers
signal.signal(signal.SIGINT, handle_shutdown)
signal.signal(signal.SIGTERM, handle_shutdown)

@app.on_event("startup")
async def startup_event():
    """Perform startup tasks"""
    logger.info("Starting up application...")
    # Ensure storage directories exist
    for dir_path in storage_dirs:
        Path(dir_path).mkdir(parents=True, exist_ok=True)
        logger.info(f"Ensured directory exists: {dir_path}")

@app.on_event("shutdown")
async def shutdown_event():
    """Perform cleanup on shutdown"""
    logger.info("Shutting down application...")
    # Add any cleanup tasks here

if __name__ == "__main__":
    try:
        port = int(os.getenv("PORT", 8000))
        logger.info(f"Starting server on port {port}")
        uvicorn.run(
            "main:app",
            host="0.0.0.0",
            port=port,
            reload=True,
            reload_excludes=["*.log", "**/logs/*", "backend/logs/*", "storage/cache/*"],
            log_level="info",
            access_log=True
        )
    except Exception as e:
        logger.error(f"Error starting server: {e}", exc_info=True)
        sys.exit(1)